import xlwings as xw
import time
from docx import Document
import shutil
import os

"""
ToDo: 
Studyname Matching
Multi Sheets in one kitBuild
Docx format
OrderLot format 02072025KM01
"""
# Copy templates to working files before editing
script_dir = os.path.dirname(os.path.abspath(__file__))
# ----------------------------------------------------------------
# Read from PartsTracker and PartsTracker2
# Read and print all sheet names in PartsTracker.xlsx
print('----------------------------------------------------------------')
partsTrackerPath = os.path.join(script_dir, 'PartsTracker.xlsx')
if os.path.exists(partsTrackerPath):
    partsWB = xw.Book(partsTrackerPath)
    # Open the sheet named '2025 General Inventory Tracker' and print all its rows
    if '2025 General Inventory Tracker' in [sheet.name for sheet in partsWB.sheets]:
        sheet = partsWB.sheets['2025 General Inventory Tracker']
        print(f"Reading sheet: {sheet.name}")
        used_range = sheet.used_range
        values = used_range.value
        for row in values:
            print(row)
    else:
        print("Sheet '2025 General Inventory Tracker' not found.")
    partsWB.close()
print('----------------------------------------------------------------')
partsTrackerPath = os.path.join(script_dir, 'PartsTracker2.xlsx')
if os.path.exists(partsTrackerPath):
    partsWB = xw.Book(partsTrackerPath)
    # Open the sheet named '2025 General Inventory Tracker' and print all its rows
    if 'Study Specific Parts Tracker' in [sheet.name for sheet in partsWB.sheets]:
        sheet = partsWB.sheets['Study Specific Parts Tracker']
        print(f"Reading sheet: {sheet.name}")
        used_range = sheet.used_range
        values = used_range.value
        for row in values:
            print(row)
    else:
        print("Sheet 'Study Specific Parts Tracker' not found.")
    partsWB.close()
print('----------------------------------------------------------------')
"""
# ----------------------------------------------------------------
def copy_if_needed(src, dst):
    if not os.path.exists(dst):
        shutil.copy(src, dst)

# Excel templates
kitBuildTemplate = os.path.join(script_dir, 'kitBuildTemplate.xlsx')
kitBuildTest = os.path.join(script_dir, 'kitBuildTest.xlsx')

pickListTemplate = os.path.join(script_dir, 'PickListTemplate.xlsx')
pickListTest = os.path.join(script_dir, 'PickListTest.xlsx')

kitQCTemplate = os.path.join(script_dir, 'KITQCTemplate.xlsx')
kitQCTest = os.path.join(script_dir, 'KITQCTest.xlsx')

# DOCX template
LabelsTemplate = os.path.join(script_dir, 'LabelsTemplate.docx')
LabelsTest = os.path.join(script_dir, 'LabelsTest.docx')

copy_if_needed(kitBuildTemplate, kitBuildTest)
copy_if_needed(pickListTemplate, pickListTest)
copy_if_needed(kitQCTemplate, kitQCTest)
copy_if_needed(LabelsTemplate, LabelsTest)
# ----------------------------------------------------------------
# Read from Kit Build file
kitBuildWB = xw.Book(kitBuildTest)
kitBuildWS = kitBuildWB.sheets[0]

# Read values from specific cells
orderDate = kitBuildWS.range('A4').value  # Row 4, Col 1 (A4)
print(orderDate)
requestByDate = kitBuildWS.range('B4').value
print(requestByDate)
studyName = if two field then top is name, if not then anything before - is name
studyProtocal = kitBuildWS.range('C4').value
print(studyProtocal)
PMRequesting = kitBuildWS.range('D4').value
print(PMRequesting)
KITID = kitBuildWS.range('E4').value
print(KITID)
numKit = kitBuildWS.range('F4').value
print(numKit)
storageLocation = 'LAB 116'
kitBuildWS.range('H4').value = storageLocation
print()
print()
print()
kitParts = []
row = 7
while True:
    if kitBuildWS.range(f'A{row}').value is None:
        break
    row_values = kitBuildWS.range(f'A{row}:F{row}').value 
    kitBuildWS.range(f'I{row}').value = numKit * kitBuildWS.range(f'C{row}').value
    kitParts.append(row_values)
    print(row_values)
    row += 1
print()
print()
print()

# Modify Pick List file
pickListWB = xw.Book(pickListTest)
pickListWS = pickListWB.sheets[0]

# Write values to specific cells 
pickListWS.range('E2').value = '02072025KM01'  # Example
pickListWS.range('E3').value = studyProtocal
pickListWS.range('E4').value = KITID

# Print for verification
for row in range(1, 13):
    print(pickListWS.range(f'A{row}:F{row}').value)



for i in range(len(kitParts)-1):
    # Insert a new row at row 7
    pickListWS.range("7:7").insert('down')
    # Copy the original template row (now at row 8) to the new row 7
    pickListWS.range("A8:F8").copy(pickListWS.range("A7"))


total = 0
for row in range(0, len(kitParts)):
    print(kitParts[row][0])
    pickListWS.range(f'B{row+7}').value = kitParts[row][0]
    pickListWS.range(f'C{row+7}').value = storageLocation
    pickListWS.range(f'D{row+7}').value = 'N/A'
    quantity = kitParts[row][2] if kitParts[row][2] is not None else 0
    total += float(quantity) * numKit
    pickListWS.range(f'E{row+7}').value = float(quantity) * numKit

pickListWS.range(f'E{len(kitParts)+7}').value = total

# ----------------------------------------------------------------
# Read from Kit Build file
kitBuildWB = xw.Book(kitBuildTest)
kitBuildWS = kitBuildWB.sheets[0]

# Read values from specific cells 
orderDate = kitBuildWS.range('A4').value  # Row 4, Col 1 (A4)
print(orderDate)
requestByDate = kitBuildWS.range('B4').value
print(requestByDate)
studyProtocal = kitBuildWS.range('C4').value
print(studyProtocal)
PMRequesting = kitBuildWS.range('D4').value
print(PMRequesting)
KITID = kitBuildWS.range('E4').value
print(KITID)
numKit = kitBuildWS.range('F4').value
print(numKit)
storageLocation = 'LAB 116'
kitBuildWS.range('H4').value = storageLocation
print()
print()
print()
kitParts = []
row = 7
while True:
    if kitBuildWS.range(f'A{row}').value is None:
        break
    row_values = kitBuildWS.range(f'A{row}:F{row}').value 
    kitBuildWS.range(f'I{row}').value = numKit * kitBuildWS.range(f'C{row}').value
    kitParts.append(row_values)
    print(row_values)
    row += 1
print()
print()
print()
# ----------------------------------------------------------------
# Modify KIT QC file
KITQCWB = xw.Book(kitQCTest)
KITQCWS = KITQCWB.sheets[0]

# Write values to specific cells
KITQCWS.range('E2').value = studyName 
KITQCWS.range('E2').color = (0, 0, 255)
KITQCWS.range('E3').value = studyProtocal
KITQCWS.range('E3').color = (0, 0, 255)
KITQCWS.range('E4').value = KITID

# Print for verification
for row in range(1, 13):
    print(KITQCWS.range(f'A{row}:F{row}').value)

for i in range(len(kitParts)-1):
    # Insert a new row at row 7
    KITQCWS.range("7:7").insert('down')
    # Copy the original template row (now at row 8) to the new row 7
    KITQCWS.range("A8:F8").copy(KITQCWS.range("A7"))


total = 0
for row in range(0, len(kitParts)):
    print(kitParts[row][0])
    KITQCWS.range(f'A{row+7}').value = kitParts[row][0]
    KITQCWS.range(f'B{row+7}').value = kitParts[row][4]
    KITQCWS.range(f'C{row+7}').value = kitParts[row][5]
    quantity = kitParts[row][2] if kitParts[row][2] is not None else 0
    total += float(quantity) * numKit
    KITQCWS.range(f'D{row+7}').value = float(quantity) * numKit

KITQCWS.range(f'D{len(kitParts)+7}').value = total

# ----------------------------------------------------------------
# Modify Label Document
labelDoc = Document(LabelsTest)

# Replace STUDYPROTOCAL with actual study protocol
for paragraph in labelDoc.paragraphs:
    if 'STUDYPROTOCAL' in paragraph.text:
        paragraph.text = paragraph.text.replace('STUDYPROTOCAL', studyProtocal)

# Replace KITID with actual kit ID
for paragraph in labelDoc.paragraphs:
    if 'KITID' in paragraph.text:
        paragraph.text = paragraph.text.replace('KITID', KITID)

# Replace the •\tITEM line with the first kit part, and insert the rest
for idx, paragraph in enumerate(labelDoc.paragraphs):
    if '•\tITEM' in paragraph.text or '• ITEM' in paragraph.text or '•\u0009ITEM' in paragraph.text:
        # Replace 'ITEM' with the first kit part name
        for run in paragraph.runs:
            if 'ITEM' in run.text:
                run.text = run.text.replace('ITEM', str(kitParts[0][0]))
        # Insert the rest of the kit parts
        for j, part in enumerate(kitParts[1:], 1):
            new_paragraph = labelDoc.add_paragraph()
            new_paragraph.text = f"\t\t•\t{part[0]}"
            # Insert after the original bullet
            labelDoc._body._body.insert(idx + j, new_paragraph._element)

# Replace LOT # and NUM with actual values
for paragraph in labelDoc.paragraphs:
    
    if 'NUM' in paragraph.text:
        # Replace NUM with the number of kits
        paragraph.text = paragraph.text.replace('NUM', '02072025KM01')

# Replace KITEXPDATE with actual expiration date
for paragraph in labelDoc.paragraphs:
    if 'EXPDATE' in paragraph.text:
        # Use the first kit part's expiration date
        exp_date = kitParts[0][5] if kitParts[0][5] != 'N/A' else 'N/A'
        paragraph.text = paragraph.text.replace('EXPDATE', exp_date)
# ----------------------------------------------------------------

# Save as new files
labelDoc.save(LabelsTest)
pickListWB.save(pickListTest)
kitBuildWB.save(kitBuildTest)
KITQCWB.save(kitQCTest)
KITQCWB.close()
pickListWB.close()
kitBuildWB.close()
"""